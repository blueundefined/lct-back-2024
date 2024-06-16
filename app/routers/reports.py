from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Inches
from io import BytesIO
import tempfile

from fastapi import APIRouter, Depends, Query
from sqlalchemy.ext.asyncio import AsyncSession
from typing import List
from datetime import datetime

from app.config import config
from app.database import get_session
from .shape_func import ShapeService, ShapeGet

import os
from tempfile import NamedTemporaryFile

router = APIRouter(prefix=config.BACKEND_PREFIX)

# Assuming you have imported ShapeService and other necessary modules as in your original code
def generate_docx_for_favorite_shapes(shapes: List[ShapeGet], filename) -> str:
    doc = Document()

    # Adding a title
    doc.add_heading('Отчёт по избранным контурам', level=1)

    for shape in shapes:
        # Add shape details to the document
        doc.add_paragraph(f"Shape ID: {shape.shape_id}")
        doc.add_paragraph(f"Архивный ID обработки: {shape.shape_version}")
        doc.add_paragraph(f"Комментарий: {shape.comment}")
        doc.add_paragraph(f"Заключение: {shape.ai_gen_comment}")

        # Add more details as needed (e.g., image, center point, Yandex panorama link)

        doc.add_page_break()

    # Save the document to a temporary file
    tmp_file = NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_file.close()  # Close to allow writing

    doc.save(tmp_file.name)

    return tmp_file.name  # Return the path to the generated file

@router.get('/shapes/favorite/docx')
async def download_favorite_shapes_docx(
    db: AsyncSession = Depends(get_session),
    limit: int = Query(100, ge=1, le=1000),
    offset: int = Query(0, ge=0),
):
    # Retrieve favorite shapes from the database using ShapeService
    favorite_shapes = await ShapeService.get_all_favorite(db=db, limit=limit, offset=offset)

    # Transform result to ShapeGet models
    favorite_shapes = [ShapeGet(shape_id=row.shape_id, shape_version=row.shape_version,
                               comment=row.comment, added_to_favorites=row.added_to_favorites,
                               created_at=str(row.created_at), updated_at=str(row.updated_at),
                               ai_gen_comment=row.ai_gen_comment) for row in favorite_shapes]

    # Generate DOCX file content and get the file path
    docx_file_path = generate_docx_for_favorite_shapes(favorite_shapes)

    # Return the generated DOCX file as a response.
    return FileResponse(BytesIO(docx_content), filename=f'отчёт_по_избранным_контурам_от_{datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}.docx', media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')